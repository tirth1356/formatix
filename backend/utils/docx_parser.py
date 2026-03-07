"""Parse DOCX files and return raw text + structural info."""
from pathlib import Path
import re
from docx import Document
from docx.oxml.ns import qn
from docx.table import Table
from docx.text.paragraph import Paragraph


# Patterns to detect heading-like text
_HEADING_PATTERN = re.compile(
    r'^(?:'
    r'[IVX]+\.\s+(.+)|'  # Roman numerals: I. INTRO, II. METHODS
    r'\d+\.\s+(.+)|'       # Arabic numerals: 1. INTRO, 2. METHODS
    r'(?:Section|Chapter)\s+[IVX\d]+[.:]?\s+(.+)'  # Section I: ..., Chapter 1: ...
    r')$',
    re.IGNORECASE | re.UNICODE
)

# Common academic section keywords
_ACADEMIC_SECTIONS = {
    'abstract', 'introduction', 'background', 'related work', 'literature', 'review',
    'methodology', 'method', 'approach', 'system', 'design', 'implementation',
    'results', 'findings', 'evaluation', 'experiments', 'performance',
    'discussion', 'analysis', 'conclusion', 'conclusions', 'future work',
    'acknowledgment', 'acknowledgments', 'references', 'bibliography', 'appendix',
    'applications', 'ethical', 'considerations', 'limitations', 'contributions',
}

def _is_orphan_heading(text: str) -> tuple:
    """
    Detect plain text that looks like an orphan heading
    (styled as body text but resembles a section heading).
    
    Returns (is_orphan_heading, cleaned_text)
    """
    if not text or len(text) > 200:
        return False, ""
    
    # All caps, short, and under 100 chars → likely a heading
    if text.isupper() and len(text) < 100:
        # Check if it contains section keywords
        text_lower = text.lower()
        if any(keyword in text_lower for keyword in _ACADEMIC_SECTIONS):
            return True, text
    
    # Title case, contains academic keywords, short
    if text.istitle() and len(text) < 100:
        text_lower = text.lower()
        if any(keyword in text_lower for keyword in _ACADEMIC_SECTIONS):
            # Make sure it's not just one sentencelike (ends with period)
            if not text.endswith('.'):
                return True, text
    
    return False, ""

def _is_heading_like(text: str, para: Paragraph = None) -> tuple:
    """
    Detect if text is heading-like.
    Returns (is_heading, heading_text, level)
    """
    if not text:
        return False, "", 0
    
    # Check pattern: "I. INTRODUCTION", "II. BACKGROUND", etc.
    match = _HEADING_PATTERN.match(text)
    if match:
        # Extract the heading text from whichever group matched
        heading = next(g for g in match.groups() if g)
        return True, heading, 1
    
    # Check if paragraph is bold (often used for headings)
    if para and para.runs:
        is_bold = any(run.bold for run in para.runs)
        # If bold, all caps or title case, AND under 100 chars, likely a heading
        if is_bold and len(text) < 100 and (text.isupper() or text.istitle()):
            return True, text, 1
    
    return False, "", 0


def parse_docx(file_path: str) -> dict:
    """
    Extract raw text and structural information from a DOCX file.
    Returns dict with type, text, pages (estimated), raw_sections.
    
    Detects both formal Heading styles, heading-like text patterns,
    and orphan headings (unmarked section boundaries that look like headings).
    """
    path = Path(file_path)
    if not path.exists():
        raise FileNotFoundError(f"DOCX not found: {file_path}")

    doc = Document(file_path)
    full_text_parts = []
    raw_sections = []
    current_section = {"heading": "", "level": 0, "content": ""}

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            # Empty paragraph - might be a section separator
            if current_section.get("content"):
                raw_sections.append(current_section)
                current_section = {"heading": "", "level": 0, "content": ""}
            continue

        # Check for formal Heading styles first
        style_name = para.style.name if para.style else ""
        is_formal_heading = "heading" in style_name.lower()
        
        level = 0
        if is_formal_heading:
            try:
                level = int(style_name.replace("Heading", "").strip() or "1")
            except ValueError:
                level = 1
        else:
            # Check for heading-like patterns (roman numerals, bold, etc.)
            is_heading_pattern, heading_text, pattern_level = _is_heading_like(text, para)
            if is_heading_pattern:
                is_formal_heading = True
                text = heading_text  # Use cleaned heading text
                level = pattern_level
            else:
                # Check for "orphan heading" - looks like a section heading but isn't styled
                is_orphan_heading, orphan_text = _is_orphan_heading(text)
                if is_orphan_heading:
                    is_formal_heading = True
                    text = orphan_text
                    level = 1
        
        if is_formal_heading:
            # Save previous section and start new one
            if current_section.get("content") or current_section.get("heading"):
                raw_sections.append(current_section)
            current_section = {"heading": text, "level": level, "content": ""}
        else:
            # Add to current section content
            if current_section["heading"] or current_section["content"]:
                current_section["content"] = current_section["content"] + "\n" + text if current_section["content"] else text
            else:
                current_section = {"heading": "", "level": 0, "content": text}

        full_text_parts.append(text)

    for table in doc.tables:
        table_text = _table_to_text(table)
        full_text_parts.append(table_text)
        raw_sections.append({"heading": "[Table]", "level": 0, "content": table_text})

    if current_section.get("content") or current_section.get("heading"):
        raw_sections.append(current_section)

    full_text = "\n\n".join(full_text_parts)
    # Rough page estimate: ~3000 chars per page
    pages = max(1, len(full_text) // 3000)

    return {
        "type": "parsed_document",
        "text": full_text,
        "pages": pages,
        "raw_sections": raw_sections,
    }


def _table_to_text(table: Table) -> str:
    rows = []
    for row in table.rows:
        cells = [cell.text.strip() for cell in row.cells]
        rows.append(" | ".join(cells))
    return "\n".join(rows)
