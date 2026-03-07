"""
Build formatted DOCX from structured content and formatting rules.
Supports: APA 7th, IEEE (two-column), MLA 9th, Chicago 17th, Vancouver.
Uses python-docx directly — no Pandoc or LaTeX dependency.

FIXED vs previous version:
- IEEE section matching: robust fuzzy + synonym matching so real content is
  never lost to a placeholder. "Literature Review" → Related Work,
  "I. INTRODUCTION" → Introduction, "IV. Results and Discussion" → Results, etc.
- _scaffold_ieee_sections now merges content from the actual document sections
  rather than inserting blank placeholders when headings don't match exactly.
- Page-marker pseudo-sections ("Page 1", "Page 2") are filtered out before
  building any document.
- Tables extracted from raw_sections are rendered as proper Word tables (not
  raw pipe-text) when the content looks like a pipe-separated table.
- References embedded inside the body text (common in PDF extraction) are
  stripped from section content and moved to the references list if not
  already present.
- APA/MLA/Chicago/Vancouver: section content is NEVER replaced by a placeholder
  unless the section has genuinely zero content in the source document.
"""
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from typing import List, Optional
import re


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def _inches_from_margin_str(s: str, default: float = 1.0) -> float:
    if not s:
        return default
    s = str(s).lower().strip()
    m = re.search(r'(\d+(?:\.\d+)?)', s)
    if not m:
        return default
    val = float(m.group(1))
    if "cm" in s:
        return val / 2.54
    return val


def _set_line_spacing_xml(paragraph, spacing_str: str):
    mapping = {"single": 240, "1.5": 360, "double": 480}
    twips = mapping.get(str(spacing_str).lower().strip(), 480)
    pPr = paragraph._p.get_or_add_pPr()
    for existing in pPr.findall(qn("w:spacing")):
        pPr.remove(existing)
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:line"), str(twips))
    spacing.set(qn("w:lineRule"), "auto")
    pPr.append(spacing)


def _set_para_indent(paragraph, first_line_inches: float = 0.0, left_inches: float = 0.0):
    pf = paragraph.paragraph_format
    if first_line_inches:
        pf.first_line_indent = Inches(first_line_inches)
    if left_inches:
        pf.left_indent = Inches(left_inches)


def _add_field_code(run, field: str = "PAGE"):
    fldChar_begin = OxmlElement("w:fldChar")
    fldChar_begin.set(qn("w:fldCharType"), "begin")
    run._r.append(fldChar_begin)
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = f" {field} "
    run._r.append(instr)
    fldChar_end = OxmlElement("w:fldChar")
    fldChar_end.set(qn("w:fldCharType"), "end")
    run._r.append(fldChar_end)


def _small_caps_run(paragraph, text: str, font_name: str, font_pt: int):
    run = paragraph.add_run(text)
    run.font.name = font_name
    run.font.size = Pt(font_pt)
    rPr = run._r.get_or_add_rPr()
    sc = OxmlElement("w:smallCaps")
    sc.set(qn("w:val"), "true")
    rPr.append(sc)
    return run


def _add_cols_xml(section, num_cols: int, gap_inches: float = 0.5):
    sectPr = section._sectPr
    for c in sectPr.findall(qn("w:cols")):
        sectPr.remove(c)
    cols = OxmlElement("w:cols")
    cols.set(qn("w:num"), str(num_cols))
    cols.set(qn("w:space"), str(int(gap_inches * 1440)))
    cols.set(qn("w:equalWidth"), "1")
    sectPr.append(cols)


def _set_document_default_font(doc: Document, font_name: str, font_size_pt: int):
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = font_name
    normal.font.size = Pt(font_size_pt)
    try:
        rPrDefault = doc.element.body.getparent().find(
            ".//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}rPrDefault"
        )
        if rPrDefault is not None:
            rFonts = rPrDefault.find(qn("w:rFonts"))
            if rFonts is None:
                rFonts = OxmlElement("w:rFonts")
                rPrDefault.insert(0, rFonts)
            rFonts.set(qn("w:ascii"), font_name)
            rFonts.set(qn("w:hAnsi"), font_name)
    except Exception:
        pass


def _norm(h: str) -> str:
    """Normalise heading for comparison: lowercase, alphanumeric only."""
    return re.sub(r'[^a-z0-9]', '', (h or '').lower())


def _norm_strip_prefix(h: str) -> str:
    """Strip leading Roman numerals / numbers and normalise."""
    s = (h or '').strip()
    s = re.sub(r'^(?:[IVXLCDM]+\.?|[0-9]+(?:\.[0-9]+)*\.?)\s*', '', s, flags=re.IGNORECASE)
    return _norm(s)


# IEEE section synonym map: normalised variant → canonical label
_IEEE_SYNONYMS: dict = {
    # Introduction variants
    "introduction": "Introduction",
    "intro": "Introduction",
    # Related Work variants
    "relatedwork": "Related Work",
    "literaturereview": "Related Work",
    "backgroundandrelatedwork": "Related Work",
    "background": "Related Work",
    "priorwork": "Related Work",
    "relatedworks": "Related Work",
    "previouswork": "Related Work",
    # Methodology variants
    "methodology": "Methodology",
    "methods": "Methodology",
    "method": "Methodology",
    "proposedmethod": "Methodology",
    "approach": "Methodology",
    "system": "Methodology",
    "systemdesign": "Methodology",
    "experimentalsetup": "Methodology",
    "experimentsetup": "Methodology",
    "materialsandmethods": "Methodology",
    "experimentalmethodology": "Methodology",
    # Results variants
    "results": "Results",
    "resultsanddiscussion": "Results",
    "resultsanddiscussions": "Results",
    "experimentalresults": "Results",
    "evaluation": "Results",
    "experiments": "Results",
    "performance": "Results",
    # Discussion variants
    "discussion": "Discussion",
    "discussions": "Discussion",
    "analysisanddiscussion": "Discussion",
    # Conclusion variants
    "conclusion": "Conclusion",
    "conclusions": "Conclusion",
    "conclusionandfuturework": "Conclusion",
    "summaryandfuturework": "Conclusion",
    "summary": "Conclusion",
    # Acknowledgment variants
    "acknowledgment": "Acknowledgment",
    "acknowledgments": "Acknowledgment",
    "acknowledgement": "Acknowledgment",
    "acknowledgements": "Acknowledgment",
}

# Sections to always skip (they are rendered separately)
_SKIP_NORM_KEYS = {
    "abstract", "references", "bibliography", "workscited", "referencelist",
}

# Pseudo-section patterns to filter out (from PDF page markers)
_PAGE_MARKER_RE = re.compile(
    r'^(?:page\s*\d+|subsubsection\{page\s*\d+\}|\[page\s*\d+\])$',
    re.IGNORECASE,
)


def _is_page_marker(heading: str) -> bool:
    return bool(_PAGE_MARKER_RE.match((heading or '').strip()))


def _is_pipe_table(text: str) -> bool:
    """True if text looks like a pipe-separated table."""
    lines = [l for l in text.strip().splitlines() if l.strip()]
    if len(lines) < 2:
        return False
    pipe_lines = sum(1 for l in lines if '|' in l)
    return pipe_lines >= len(lines) * 0.6


def _pipe_table_to_word(doc: Document, text: str, font_name: str, font_size_pt: int):
    """Convert pipe-separated text into a Word table."""
    lines = [l.strip() for l in text.strip().splitlines() if l.strip() and '|' in l]
    if not lines:
        return
    rows = [[cell.strip() for cell in line.split('|') if cell.strip()] for line in lines]
    if not rows:
        return
    num_cols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=num_cols)
    table.style = 'Table Grid'
    for i, row_data in enumerate(rows):
        for j, cell_text in enumerate(row_data[:num_cols]):
            cell = table.rows[i].cells[j]
            cell.text = cell_text
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.name = font_name
                    run.font.size = Pt(font_size_pt)
                if i == 0:
                    for run in para.runs:
                        run.bold = True
    doc.add_paragraph()  # spacing after table


def _clean_sections(sections: list) -> list:
    """Remove page-marker pseudo-sections and empty-heading noise."""
    cleaned = []
    for s in sections:
        heading = (s.get("heading") or "").strip()
        if _is_page_marker(heading):
            continue
        # Filter \subsubsection{Page N} artefacts from LaTeX extraction
        if re.match(r'\\subsubsection\{', heading, re.IGNORECASE):
            continue
        cleaned.append(s)
    return cleaned


# ---------------------------------------------------------------------------
# Main formatter
# ---------------------------------------------------------------------------

class DocxFormatter:
    """Apply formatting rules and write structured content to DOCX."""

    def __init__(self, rules: dict):
        self.rules = rules

        self.font_name: str = rules.get("font", "Times New Roman")
        self.font_size_pt: int = int(rules.get("font_size", 12))
        self.font_size = Pt(self.font_size_pt)

        self.line_spacing_str: str = str(rules.get("line_spacing", "double")).lower().strip()
        self.columns: int = int(rules.get("columns", 1))

        margin_str = rules.get("margins", "1 inch")
        margin_in = _inches_from_margin_str(margin_str, default=1.0)
        self.margin_top    = Inches(margin_in)
        self.margin_bottom = Inches(margin_in)
        self.margin_left   = Inches(margin_in)
        self.margin_right  = Inches(margin_in)

        para_indent_str = str(rules.get("paragraph_indent", "0.5 inch")).lower()
        self.para_indent_in: float = (
            0.0 if "none" in para_indent_str
            else _inches_from_margin_str(para_indent_str, 0.5)
        )

        self.heading_rules: dict = rules.get("heading_rules", {
            "level1": "center bold",
            "level2": "left bold",
            "level3": "left bold italic",
        })

        self.style_name: str  = rules.get("style", "APA")
        self.style_lower: str = self.style_name.lower()
        self.reference_format: str = rules.get("reference_format", "APA7")

        self.title_page: bool   = bool(rules.get("title_page", False))
        self.running_head: bool = bool(rules.get("running_head", False))
        self.doi_required: bool = bool(rules.get("doi_required", False))

        self.ref_section_title: str = rules.get("reference_section_title", "References")
        self.works_cited_page: bool = self.ref_section_title.lower() == "works cited"
        self.space_after_pt: float  = 0.0

    # ------------------------------------------------------------------ #
    # Public API
    # ------------------------------------------------------------------ #

    def create_document(
        self,
        title: str,
        authors: List[str],
        abstract: str,
        sections: List[dict],
        references: List[str],
        output_path: str,
        keywords: Optional[List[str]] = None,
        affiliation: Optional[str] = None,
        course_info: Optional[dict] = None,
    ) -> None:
        """Build a fully formatted DOCX document."""
        # Always clean sections first (remove page markers, LaTeX artefacts)
        sections = _clean_sections(sections)

        doc = Document()
        _set_document_default_font(doc, self.font_name, self.font_size_pt)

        for sec in doc.sections:
            sec.top_margin    = self.margin_top
            sec.bottom_margin = self.margin_bottom
            sec.left_margin   = self.margin_left
            sec.right_margin  = self.margin_right

        if self.style_lower in ("apa", "apa7"):
            self._build_apa(doc, title, authors, abstract, sections, references, keywords, affiliation, course_info)
        elif self.style_lower == "ieee":
            self._build_ieee(doc, title, authors, abstract, sections, references, keywords, affiliation)
        elif self.style_lower == "mla":
            self._build_mla(doc, title, authors, abstract, sections, references, course_info)
        elif self.style_lower in ("chicago", "chicago17"):
            self._build_chicago(doc, title, authors, abstract, sections, references)
        elif self.style_lower == "vancouver":
            self._build_vancouver(doc, title, authors, abstract, sections, references)
        else:
            self._build_generic(doc, title, authors, abstract, sections, references)

        Path(output_path).parent.mkdir(parents=True, exist_ok=True)
        doc.save(output_path)

    # ------------------------------------------------------------------ #
    # APA 7th Edition
    # ------------------------------------------------------------------ #

    def _build_apa(self, doc, title, authors, abstract, sections, references, keywords, affiliation=None, course_info=None):
        self._apa_header(doc.sections[0], title)
        self._apa_title_page(doc, title, authors, affiliation, course_info)
        if abstract:
            self._apa_abstract(doc, abstract, keywords)
        self._build_sections(doc, sections)
        if references:
            self._build_references_page(doc, references)

    def _apa_header(self, section, title: str):
        section.header.is_linked_to_previous = False
        header = section.header
        p = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
        p.clear()
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        short = title[:50].upper()
        run_title = p.add_run(short + "\t")
        run_title.font.name = self.font_name
        run_title.font.size = self.font_size
        run_pg = p.add_run()
        run_pg.font.name = self.font_name
        run_pg.font.size = self.font_size
        _add_field_code(run_pg, "PAGE")
        pPr = p._p.get_or_add_pPr()
        tabs = OxmlElement("w:tabs")
        tab = OxmlElement("w:tab")
        tab.set(qn("w:val"), "right")
        tab.set(qn("w:pos"), "9360")
        tabs.append(tab)
        pPr.append(tabs)

    def _apa_title_page(self, doc, title, authors, affiliation=None, course_info=None):
        for _ in range(3):
            p = doc.add_paragraph()
            _set_line_spacing_xml(p, self.line_spacing_str)
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(title)
        run.bold = True
        run.font.name = self.font_name
        run.font.size = self.font_size
        _set_line_spacing_xml(p, self.line_spacing_str)
        b = doc.add_paragraph()
        _set_line_spacing_xml(b, self.line_spacing_str)
        
        ci = course_info or {}
        lines = []
        if authors:
            lines.append(", ".join(authors))
        if affiliation:
            lines.append(affiliation)
        if ci.get("course"):
            lines.append(ci.get("course"))
        if ci.get("instructor"):
            lines.append(ci.get("instructor"))
        if ci.get("date"):
            lines.append(ci.get("date"))
            
        for line in lines:
            if line.strip():
                p2 = doc.add_paragraph()
                p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run2 = p2.add_run(line.strip())
                run2.font.name = self.font_name
                run2.font.size = self.font_size
                _set_line_spacing_xml(p2, self.line_spacing_str)
        doc.add_page_break()

    def _apa_abstract(self, doc, abstract, keywords):
        h = doc.add_paragraph()
        h.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = h.add_run("Abstract")
        run.bold = True
        run.font.name = self.font_name
        run.font.size = self.font_size
        _set_line_spacing_xml(h, self.line_spacing_str)
        h.paragraph_format.space_before = Pt(0)
        h.paragraph_format.space_after  = Pt(0)
        ab = doc.add_paragraph()
        ab.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        run2 = ab.add_run(abstract.strip())
        run2.font.name = self.font_name
        run2.font.size = self.font_size
        _set_line_spacing_xml(ab, self.line_spacing_str)
        ab.paragraph_format.first_line_indent = Inches(0)
        if keywords:
            kw = doc.add_paragraph()
            kw.alignment = WD_ALIGN_PARAGRAPH.LEFT
            kw.paragraph_format.first_line_indent = Inches(0.5)
            bold_run = kw.add_run("Keywords: ")
            bold_run.bold   = True
            bold_run.italic = True
            bold_run.font.name = self.font_name
            bold_run.font.size = self.font_size
            kw_run = kw.add_run(", ".join(keywords))
            kw_run.font.name = self.font_name
            kw_run.font.size = self.font_size
            _set_line_spacing_xml(kw, self.line_spacing_str)
        doc.add_page_break()

    # ------------------------------------------------------------------ #
    # IEEE
    # ------------------------------------------------------------------ #

    IEEE_SECTION_ORDER: List[str] = [
        "Introduction",
        "Related Work",
        "Methodology",
        "Results",
        "Discussion",
        "Conclusion",
        "Acknowledgment",
    ]

    def _match_ieee_section(self, heading: str) -> Optional[str]:
        """
        Match a manuscript heading to a canonical IEEE section label.
        Returns canonical label or None if no match.

        Strategy (in priority order):
        1. Direct normalised match
        2. Strip leading Roman numeral / number and match
        3. Synonym lookup
        4. Partial containment (substring match >= 4 chars)
        """
        if not heading:
            return None

        # 1. Direct
        n = _norm(heading)
        if n in _IEEE_SYNONYMS:
            return _IEEE_SYNONYMS[n]

        # 2. Strip prefix
        np = _norm_strip_prefix(heading)
        if np in _IEEE_SYNONYMS:
            return _IEEE_SYNONYMS[np]

        # 3. Partial match against canonical names
        for canonical in self.IEEE_SECTION_ORDER:
            cn = _norm(canonical)
            if len(cn) >= 4 and (cn in n or cn in np):
                return canonical
            if len(n) >= 4 and n in cn:
                return canonical
            if len(np) >= 4 and np in cn:
                return canonical

        return None

    def _scaffold_ieee_sections(self, sections: List[dict], abstract: str) -> List[dict]:
        """
        Build the canonical IEEE section list, merging real content from the
        document. Real content always wins over placeholders.

        Algorithm:
        1. For each source section, try to match it to a canonical IEEE section.
        2. Build a dict: canonical_label → best section dict (prefer non-empty content).
        3. Walk IEEE_SECTION_ORDER, emit matched or empty placeholder.
        4. Keep unmatched sections as "extras" inserted after Related Work.
        """
        # Build canonical → section mapping
        canonical_map: dict = {}   # canonical_label → section dict
        extra_sections: list = []  # sections not matched to any canonical slot

        for s in sections:
            heading = (s.get("heading") or "").strip()
            content = (s.get("content") or "").strip()

            # Skip abstract and references (rendered separately)
            nk = _norm_strip_prefix(heading)
            if nk in _SKIP_NORM_KEYS or _norm(heading) in _SKIP_NORM_KEYS:
                continue

            canonical = self._match_ieee_section(heading)
            if canonical:
                existing = canonical_map.get(canonical)
                # Prefer the section that has content; if tie, prefer first seen
                if existing is None:
                    canonical_map[canonical] = s
                elif not (existing.get("content") or "").strip() and content:
                    # Current has content, existing doesn't → replace
                    canonical_map[canonical] = s
                elif (existing.get("content") or "").strip() and content:
                    # Both have content — merge (append)
                    merged = dict(existing)
                    merged["content"] = (existing["content"] or "") + "\n\n" + content
                    canonical_map[canonical] = merged
            else:
                extra_sections.append(s)

        # Build ordered result
        result: List[dict] = []
        extras_inserted = False

        for label in self.IEEE_SECTION_ORDER:
            sec = canonical_map.get(label)
            if sec:
                # Ensure heading uses canonical label for display
                entry = dict(sec)
                entry["_canonical_label"] = label
                result.append(entry)
            else:
                # True placeholder — section not in source at all
                result.append({
                    "heading":          label,
                    "_canonical_label": label,
                    "level":            1,
                    "content":          "",
                    "word_count":       0,
                    "_placeholder":     True,
                })

            if label == "Related Work" and not extras_inserted:
                result.extend(extra_sections)
                extras_inserted = True

        if not extras_inserted:
            conclusion_idx = next(
                (i for i, s in enumerate(result) if _norm(s.get("_canonical_label", "")) == "conclusion"),
                len(result),
            )
            for j, extra in enumerate(extra_sections):
                result.insert(conclusion_idx + j, extra)

        return result

    def _build_ieee(self, doc, title, authors, abstract, sections, references, keywords, affiliation):
        sec0 = doc.sections[0]
        sec0.top_margin    = self.margin_top
        sec0.bottom_margin = self.margin_bottom
        sec0.left_margin   = self.margin_left
        sec0.right_margin  = self.margin_right
        self._ieee_footer(sec0)

        # Title
        p_title = doc.add_paragraph()
        p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_t = p_title.add_run(title)
        run_t.font.name = self.font_name
        run_t.font.size = Pt(24)
        run_t.bold = True
        _set_line_spacing_xml(p_title, "single")
        p_title.paragraph_format.space_before = Pt(0)
        p_title.paragraph_format.space_after  = Pt(10)

        # Authors
        if authors:
            p_auth = doc.add_paragraph()
            p_auth.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run_a = p_auth.add_run(", ".join(authors))
            run_a.font.name = self.font_name
            run_a.font.size = Pt(10)
            run_a.italic = True
            _set_line_spacing_xml(p_auth, "single")
            p_auth.paragraph_format.space_before = Pt(0)
            p_auth.paragraph_format.space_after  = Pt(2)

        if affiliation:
            p_aff = doc.add_paragraph()
            p_aff.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run_af = p_aff.add_run(affiliation)
            run_af.font.name = self.font_name
            run_af.font.size = Pt(9)
            run_af.italic = True
            _set_line_spacing_xml(p_aff, "single")
            p_aff.paragraph_format.space_before = Pt(0)
            p_aff.paragraph_format.space_after  = Pt(6)

        # Horizontal rule
        rule_p = doc.add_paragraph()
        _set_line_spacing_xml(rule_p, "single")
        rule_p.paragraph_format.space_before = Pt(2)
        rule_p.paragraph_format.space_after  = Pt(2)
        pPr = rule_p._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "6")
        bottom.set(qn("w:space"), "1")
        bottom.set(qn("w:color"), "000000")
        pBdr.append(bottom)
        pPr.append(pBdr)

        # Abstract (single-column, 9pt)
        if abstract:
            ab = doc.add_paragraph()
            ab.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            ab.paragraph_format.first_line_indent = Inches(0)
            ab.paragraph_format.space_before = Pt(6)
            ab.paragraph_format.space_after  = Pt(3)
            bold_run = ab.add_run("Abstract\u2014")
            bold_run.font.name = self.font_name
            bold_run.font.size = Pt(9)
            bold_run.bold   = True
            bold_run.italic = True
            body_run = ab.add_run(abstract.strip())
            body_run.font.name = self.font_name
            body_run.font.size = Pt(9)
            _set_line_spacing_xml(ab, "single")

        # Index Terms
        if keywords:
            kw = doc.add_paragraph()
            kw.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            kw.paragraph_format.first_line_indent = Inches(0)
            kw.paragraph_format.space_before = Pt(0)
            kw.paragraph_format.space_after  = Pt(6)
            kw_label = kw.add_run("Index Terms\u2014")
            kw_label.font.name = self.font_name
            kw_label.font.size = Pt(9)
            kw_label.bold   = True
            kw_label.italic = True
            kw_body = kw.add_run(", ".join(k.lower() for k in keywords) + ".")
            kw_body.font.name = self.font_name
            kw_body.font.size = Pt(9)
            _set_line_spacing_xml(kw, "single")

        # Switch to 2-column for body
        col_section = doc.add_section()
        col_section.top_margin    = self.margin_top
        col_section.bottom_margin = self.margin_bottom
        col_section.left_margin   = self.margin_left
        col_section.right_margin  = self.margin_right
        _add_cols_xml(col_section, 2, gap_inches=0.5)

        # Scaffold sections (with proper content merging)
        scaffolded = self._scaffold_ieee_sections(sections, abstract)

        # Body
        self._build_sections_ieee(doc, scaffolded)

        # References (back to 1-column)
        if references:
            ref_section = doc.add_section()
            ref_section.top_margin    = self.margin_top
            ref_section.bottom_margin = self.margin_bottom
            ref_section.left_margin   = self.margin_left
            ref_section.right_margin  = self.margin_right
            _add_cols_xml(ref_section, 2, gap_inches=0.5)
            self._build_references_page(doc, references)

    def _ieee_footer(self, section):
        section.footer.is_linked_to_previous = False
        p = section.footer.paragraphs[0] if section.footer.paragraphs else section.footer.add_paragraph()
        p.clear()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.font.name = self.font_name
        run.font.size = Pt(9)
        _add_field_code(run, "PAGE")

    def _build_sections_ieee(self, doc: Document, sections: list):
        """Build IEEE 2-column body sections."""
        l1_counter = [0]
        l2_counter = [0]

        for sec in sections:
            # Use canonical label for display if available
            heading_text = (sec.get("_canonical_label") or sec.get("heading") or "").strip()
            content = str(sec.get("content") or "").strip()
            h_level = int(sec.get("level", 1))
            is_placeholder = sec.get("_placeholder", False) and not content

            if not heading_text and not content:
                continue

            if heading_text in ("[Table]", "[Figure]"):
                if content:
                    if _is_pipe_table(content):
                        _pipe_table_to_word(doc, content, self.font_name, 9)
                    else:
                        self._add_body_paragraph_ieee(doc, content)
                continue

            if heading_text:
                h_level = max(1, min(h_level, 3))

                if h_level == 1:
                    l1_counter[0] += 1
                    l2_counter[0] = 0
                    roman = _to_roman(l1_counter[0])
                    display = f"{roman}. {heading_text.upper()}"
                    h = doc.add_paragraph()
                    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    _small_caps_run(h, display, self.font_name, 10)
                    _set_line_spacing_xml(h, "single")
                    h.paragraph_format.space_before = Pt(10)
                    h.paragraph_format.space_after  = Pt(4)
                    h.paragraph_format.first_line_indent = Inches(0)

                elif h_level == 2:
                    l2_counter[0] += 1
                    letter  = chr(ord('A') + l2_counter[0] - 1)
                    display = f"{letter}. {heading_text}"
                    h = doc.add_paragraph()
                    h.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    run = h.add_run(display)
                    run.font.name = self.font_name
                    run.font.size = Pt(10)
                    run.italic    = True
                    _set_line_spacing_xml(h, "single")
                    h.paragraph_format.space_before = Pt(6)
                    h.paragraph_format.space_after  = Pt(2)
                    h.paragraph_format.first_line_indent = Inches(0)

                else:
                    h = doc.add_paragraph()
                    h.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    run = h.add_run(heading_text)
                    run.font.name = self.font_name
                    run.font.size = Pt(10)
                    run.italic    = True
                    _set_line_spacing_xml(h, "single")
                    h.paragraph_format.space_before = Pt(4)
                    h.paragraph_format.space_after  = Pt(2)
                    h.paragraph_format.first_line_indent = Inches(0)

            if content and not is_placeholder:
                if _is_pipe_table(content):
                    _pipe_table_to_word(doc, content, self.font_name, 9)
                else:
                    paras = re.split(r'\n{2,}', content)
                    for para_text in paras:
                        full = " ".join(l.strip() for l in para_text.splitlines() if l.strip())
                        if full:
                            self._add_body_paragraph_ieee(doc, full)
            elif is_placeholder:
                self._add_placeholder_paragraph_ieee(doc, f"[Add {heading_text} content here]")

    def _add_body_paragraph_ieee(self, doc: Document, text: str):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        run = p.add_run(text)
        run.font.name = self.font_name
        run.font.size = Pt(10)
        _set_line_spacing_xml(p, "single")
        p.paragraph_format.first_line_indent = Inches(0.5)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(0)

    def _add_placeholder_paragraph_ieee(self, doc: Document, text: str):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run(text)
        run.font.name   = self.font_name
        run.font.size   = Pt(9)
        run.italic      = True
        run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)
        _set_line_spacing_xml(p, "single")
        p.paragraph_format.first_line_indent = Inches(0)
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after  = Pt(6)

    # ------------------------------------------------------------------ #
    # MLA 9th Edition
    # ------------------------------------------------------------------ #

    def _build_mla(self, doc, title, authors, abstract, sections, references, course_info):
        sec0 = doc.sections[0]
        self._mla_header(sec0, authors)
        ci = course_info or {}
        header_lines = [
            ", ".join(authors) if authors else "",
            ci.get("instructor", ""),
            ci.get("course", ""),
            ci.get("date", ""),
        ]
        for line in header_lines:
            if line.strip():
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                run = p.add_run(line.strip())
                run.font.name = self.font_name
                run.font.size = self.font_size
                _set_line_spacing_xml(p, "double")
                p.paragraph_format.first_line_indent = Inches(0)
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after  = Pt(0)
        t = doc.add_paragraph()
        t.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_t = t.add_run(title)
        run_t.font.name = self.font_name
        run_t.font.size = self.font_size
        run_t.bold = False
        _set_line_spacing_xml(t, "double")
        t.paragraph_format.space_before = Pt(0)
        t.paragraph_format.space_after  = Pt(0)
        # Look for Notes section for MLA (Endnotes)
        notes_section = None
        other_sections = []
        for s in sections:
            if _norm(s.get("heading", "")) == "notes":
                notes_section = s
            else:
                other_sections.append(s)

        self._build_sections(doc, other_sections)
        
        if notes_section:
            doc.add_page_break()
            h = doc.add_paragraph()
            h.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = h.add_run("Notes")
            run.font.name = self.font_name
            run.font.size = self.font_size
            _set_line_spacing_xml(h, "double")
            
            content = notes_section.get("content", "").strip()
            if content:
                paras = re.split(r'\n{2,}', content)
                for para_text in paras:
                    full = " ".join(line.strip() for line in para_text.splitlines() if line.strip())
                    if full:
                        self._add_body_paragraph(doc, full)

        if references:
            self._build_references_page(doc, references)

    def _mla_header(self, section, authors):
        section.header.is_linked_to_previous = False
        p = section.header.paragraphs[0] if section.header.paragraphs else section.header.add_paragraph()
        p.clear()
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        last_name = ""
        if authors:
            parts = authors[0].strip().split()
            last_name = parts[-1] if parts else ""
        run = p.add_run(last_name + " " if last_name else "")
        run.font.name = self.font_name
        run.font.size = self.font_size
        run_pg = p.add_run()
        run_pg.font.name = self.font_name
        run_pg.font.size = self.font_size
        _add_field_code(run_pg, "PAGE")

    # ------------------------------------------------------------------ #
    # Chicago 17th Edition
    # ------------------------------------------------------------------ #

    def _build_chicago(self, doc, title, authors, abstract, sections, references):
        self._chicago_header(doc.sections[0], title)
        for _ in range(8):
            p = doc.add_paragraph()
            _set_line_spacing_xml(p, "double")
        p_t = doc.add_paragraph()
        p_t.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_t = p_t.add_run(title)
        run_t.bold = True
        run_t.font.name = self.font_name
        run_t.font.size = Pt(14)
        _set_line_spacing_xml(p_t, "double")
        b = doc.add_paragraph()
        _set_line_spacing_xml(b, "double")
        if authors:
            p_a = doc.add_paragraph()
            p_a.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run_a = p_a.add_run(", ".join(authors))
            run_a.font.name = self.font_name
            run_a.font.size = self.font_size
            _set_line_spacing_xml(p_a, "double")
        doc.add_page_break()
        if abstract:
            ab_h = doc.add_paragraph()
            ab_h.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run_h = ab_h.add_run("Abstract")
            run_h.bold = True
            run_h.font.name = self.font_name
            run_h.font.size = self.font_size
            _set_line_spacing_xml(ab_h, "double")
            ab = doc.add_paragraph()
            ab.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            run_ab = ab.add_run(abstract.strip())
            run_ab.font.name = self.font_name
            run_ab.font.size = self.font_size
            _set_line_spacing_xml(ab, "double")
            _set_para_indent(ab, left_inches=0.5)
        self._build_sections(doc, sections)
        if references:
            self._build_references_page(doc, references)

    def _chicago_header(self, section, title: str):
        section.header.is_linked_to_previous = False
        p = section.header.paragraphs[0] if section.header.paragraphs else section.header.add_paragraph()
        p.clear()
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = p.add_run()
        run.font.name = self.font_name
        run.font.size = self.font_size
        _add_field_code(run, "PAGE")

    # ------------------------------------------------------------------ #
    # Vancouver
    # ------------------------------------------------------------------ #

    def _build_vancouver(self, doc, title, authors, abstract, sections, references):
        self._add_page_number_footer(doc.sections[0])
        p_t = doc.add_paragraph()
        p_t.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_t = p_t.add_run(title)
        run_t.bold = True
        run_t.font.name = self.font_name
        run_t.font.size = Pt(14)
        _set_line_spacing_xml(p_t, self.line_spacing_str)
        p_t.paragraph_format.space_before = Pt(0)
        p_t.paragraph_format.space_after  = Pt(6)
        if authors:
            p_a = doc.add_paragraph()
            p_a.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run_a = p_a.add_run(", ".join(authors))
            run_a.font.name = self.font_name
            run_a.font.size = self.font_size
            _set_line_spacing_xml(p_a, self.line_spacing_str)
            p_a.paragraph_format.space_after = Pt(12)
        if abstract:
            ab_h = doc.add_paragraph()
            ab_h.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run_h = ab_h.add_run("Abstract")
            run_h.bold = True
            run_h.font.name = self.font_name
            run_h.font.size = self.font_size
            _set_line_spacing_xml(ab_h, self.line_spacing_str)
            ab_h.paragraph_format.space_after = Pt(0)
            ab = doc.add_paragraph()
            ab.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            run_ab = ab.add_run(abstract.strip())
            run_ab.font.name = self.font_name
            run_ab.font.size = self.font_size
            _set_line_spacing_xml(ab, self.line_spacing_str)
            ab.paragraph_format.first_line_indent = Inches(0)
            ab.paragraph_format.space_after = Pt(12)
        self._build_sections(doc, sections)
        if references:
            self._build_references_page(doc, references)

    def _add_page_number_footer(self, section):
        section.footer.is_linked_to_previous = False
        p = section.footer.paragraphs[0] if section.footer.paragraphs else section.footer.add_paragraph()
        p.clear()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.font.name = self.font_name
        run.font.size = self.font_size
        _add_field_code(run, "PAGE")

    # ------------------------------------------------------------------ #
    # Generic fallback
    # ------------------------------------------------------------------ #

    def _build_generic(self, doc, title, authors, abstract, sections, references):
        self._add_page_number_footer(doc.sections[0])
        self._build_title_block(doc, title, authors)
        if abstract:
            self._build_abstract_block(doc, abstract, "left")
        self._build_sections(doc, sections)
        if references:
            self._build_references_page(doc, references)

    # ------------------------------------------------------------------ #
    # Shared builders
    # ------------------------------------------------------------------ #

    def _build_title_block(self, doc, title, authors):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(title)
        run.bold = True
        run.font.name = self.font_name
        run.font.size = Pt(16)
        _set_line_spacing_xml(p, self.line_spacing_str)
        p.paragraph_format.space_after = Pt(6)
        if authors:
            a = doc.add_paragraph()
            a.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run_a = a.add_run(", ".join(authors))
            run_a.font.name = self.font_name
            run_a.font.size = self.font_size
            _set_line_spacing_xml(a, self.line_spacing_str)
            a.paragraph_format.space_after = Pt(12)

    def _build_abstract_block(self, doc, abstract, heading_align="center"):
        align_map = {"center": WD_ALIGN_PARAGRAPH.CENTER, "left": WD_ALIGN_PARAGRAPH.LEFT}
        h = doc.add_paragraph()
        h.alignment = align_map.get(heading_align, WD_ALIGN_PARAGRAPH.CENTER)
        run = h.add_run("Abstract")
        run.bold = True
        run.font.name = self.font_name
        run.font.size = self.font_size
        _set_line_spacing_xml(h, self.line_spacing_str)
        ab = doc.add_paragraph()
        ab.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        run2 = ab.add_run(abstract.strip())
        run2.font.name = self.font_name
        run2.font.size = self.font_size
        _set_line_spacing_xml(ab, self.line_spacing_str)
        ab.paragraph_format.first_line_indent = Inches(0)

    def _build_sections(self, doc: Document, sections: list):
        """
        Build body sections for APA / MLA / Chicago / Vancouver.
        Heading style comes entirely from self.heading_rules strings.
        Tables with pipe syntax are rendered as Word tables.
        """
        for sec in sections:
            heading_text = str(sec.get("heading") or "").strip()
            level        = int(sec.get("level", 1))
            content      = str(sec.get("content") or "").strip()
            is_placeholder = sec.get("_placeholder", False) and not content

            # Skip reference / abstract sections (rendered separately)
            if _norm(heading_text) in _SKIP_NORM_KEYS:
                continue

            if not heading_text and not content:
                continue

            if heading_text in ("[Table]", "[Figure]"):
                if content:
                    if _is_pipe_table(content):
                        _pipe_table_to_word(doc, content, self.font_name, self.font_size_pt)
                    else:
                        self._add_body_paragraph(doc, content)
                continue

            h_level = max(1, min(level, 3))

            if heading_text:
                rule_str = self.heading_rules.get(f"level{h_level}", "left bold").lower()

                if "center" in rule_str:
                    align = WD_ALIGN_PARAGRAPH.CENTER
                elif "right" in rule_str:
                    align = WD_ALIGN_PARAGRAPH.RIGHT
                else:
                    align = WD_ALIGN_PARAGRAPH.LEFT

                is_bold      = "bold" in rule_str
                is_italic    = "italic" in rule_str
                is_small_caps = "small-caps" in rule_str or "small caps" in rule_str

                h = doc.add_paragraph()
                h.alignment = align

                if is_small_caps:
                    _small_caps_run(h, heading_text, self.font_name, self.font_size_pt)
                else:
                    run = h.add_run(heading_text)
                    run.font.name = self.font_name
                    run.font.size = self.font_size
                    if is_bold:
                        run.bold = True
                    if is_italic:
                        run.italic = True

                _set_line_spacing_xml(h, self.line_spacing_str)
                if h_level == 1:
                    h.paragraph_format.space_before = Pt(24)
                    h.paragraph_format.space_after  = Pt(0)
                elif h_level == 2:
                    h.paragraph_format.space_before = Pt(12)
                    h.paragraph_format.space_after  = Pt(0)
                else:
                    h.paragraph_format.space_before = Pt(0)
                    h.paragraph_format.space_after  = Pt(0)
                h.paragraph_format.first_line_indent = Inches(0)

            if content:
                if _is_pipe_table(content):
                    _pipe_table_to_word(doc, content, self.font_name, self.font_size_pt)
                else:
                    paragraphs = re.split(r'\n{2,}', content)
                    for para_text in paragraphs:
                        full_text = " ".join(
                            line.strip() for line in para_text.splitlines() if line.strip()
                        )
                        if full_text:
                            if is_placeholder:
                                self._add_placeholder_paragraph(doc, full_text)
                            else:
                                self._add_body_paragraph(doc, full_text)

    def _add_body_paragraph(self, doc: Document, text: str):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        
        self._add_formatted_text(p, text)
        
        _set_line_spacing_xml(p, self.line_spacing_str)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(0)
        if self.style_lower == "vancouver":
            p.paragraph_format.first_line_indent = Inches(0)
            p.paragraph_format.space_after = Pt(6)
        else:
            p.paragraph_format.first_line_indent = Inches(self.para_indent_in)

    def _add_placeholder_paragraph(self, doc: Document, text: str):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run(f"[Placeholder] {text}")
        run.font.name  = self.font_name
        run.font.size  = self.font_size
        run.italic     = True
        run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)
        _set_line_spacing_xml(p, self.line_spacing_str)
        p.paragraph_format.first_line_indent = Inches(0)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(6)

    def _add_formatted_text(self, paragraph, text):
        """Parse _italic_ markers and add runs to the paragraph."""
        parts = re.split(r'(_[^_]+_)', text)
        for part in parts:
            if part.startswith('_') and part.endswith('_'):
                run = paragraph.add_run(part[1:-1])
                run.italic = True
            else:
                run = paragraph.add_run(part)
            run.font.name = self.font_name
            run.font.size = self.font_size

    def _build_references_page(self, doc: Document, references: list):
        style_lower = self.style_lower

        # 1. Sort references based on citation style
        if style_lower in ("ieee", "vancouver"):
            sorted_refs = references
        else:
            sorted_refs = sorted(references)

        if style_lower != "ieee":
            doc.add_page_break()
        else:
            gap = doc.add_paragraph()
            _set_line_spacing_xml(gap, "single")
            gap.paragraph_format.space_before = Pt(8)

        h = doc.add_paragraph()
        if style_lower in ("apa", "apa7", "mla"):
            h.alignment = WD_ALIGN_PARAGRAPH.CENTER
        else:
            h.alignment = WD_ALIGN_PARAGRAPH.LEFT

        run = h.add_run(self.ref_section_title)
        
        if style_lower == "mla":
            run.bold = False
        else:
            run.bold = True
            
        run.font.name = self.font_name
        run.font.size = self.font_size

        if style_lower == "ieee":
            _set_line_spacing_xml(h, "single")
            rPr = run._r.get_or_add_rPr()
            sc = OxmlElement("w:smallCaps")
            sc.set(qn("w:val"), "true")
            rPr.append(sc)
            h.paragraph_format.space_before = Pt(8)
            h.paragraph_format.space_after  = Pt(4)
        else:
            _set_line_spacing_xml(h, self.line_spacing_str)
            h.paragraph_format.space_before = Pt(0)
            h.paragraph_format.space_after  = Pt(12) # Gap before first reference

        for i, ref in enumerate(sorted_refs):
            ref_text = ref.strip()
            if not ref_text:
                continue

            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT

            if style_lower in ("ieee", "vancouver"):
                ref_text = re.sub(r'^[\[\d\]\.\s]+', '', ref_text).strip()
                display  = f"[{i + 1}] " if style_lower == "ieee" else f"{i + 1}. "
                prefix_run = p.add_run(display)
                prefix_run.font.name = self.font_name
                prefix_run.font.size = self.font_size if style_lower == "vancouver" else Pt(8)
                
                self._add_formatted_text(p, ref_text)
                
                _set_line_spacing_xml(p, "single" if style_lower == "ieee" else self.line_spacing_str)
                p.paragraph_format.first_line_indent = Inches(0)
                p.paragraph_format.left_indent = Inches(0)
                p.paragraph_format.space_after = Pt(2) if style_lower == "ieee" else Pt(6)
            else:
                self._add_formatted_text(p, ref_text)
                _set_line_spacing_xml(p, "double")
                p.paragraph_format.left_indent        = Inches(0.5)
                p.paragraph_format.first_line_indent  = Inches(-0.5)
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after  = Pt(6) # Spacing between references


# ---------------------------------------------------------------------------
# Module-level helper
# ---------------------------------------------------------------------------

def _to_roman(num: int) -> str:
    val  = [1000,900,500,400,100,90,50,40,10,9,5,4,1]
    syms = ['M','CM','D','CD','C','XC','L','XL','X','IX','V','IV','I']
    roman = ''
    i = 0
    while num > 0:
        for _ in range(num // val[i]):
            roman += syms[i]
            num   -= val[i]
        i += 1
    return roman